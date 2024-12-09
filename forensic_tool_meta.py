import PyPDF2
import re
from docx import Document
import pandas as pd
import openpyxl
from langdetect import detect
import geopy
from geopy.geocoders import Nominatim
import exifread
import sqlite3


class Forensic:

    @staticmethod
    def _convert_to_degrees(value):
        """
        Convert GPS coordinates from degrees, minutes, seconds to decimal degrees.
        :param value: The value to convert.
        :return: Decimal degrees.
        """
        d = float(value[0])
        m = float(value[1])
        s = float(value[2])
        return d + (m / 60.0) + (s / 3600.0)


    @staticmethod
    def get_gps_from_exif(img_filename):
        with open(img_filename, "rb") as file:
            exif = exifread.process_file(file)
            if not exif:
                print("NOTHING METADATA FOUND")
            else:
                latitude = exif.get("GPS GPSLatitude")
                latitude_ref = exif.get("GPS GPSLatitudeRef")
                longitude = exif.get("GPS GPSLongitude")
                longitude_ref = exif.get("GPS GPSLongitudeRef")
                altitude = exif.get("GPS GPSAltitude")
                altitude_ref = exif.get("GPS GPSAltitudeRef")

                print(str(latitude) + " " + str(longitude))
                if latitude and latitude_ref and longitude and longitude_ref:
                    # Extraction des valeurs numériques des objets IfdTag
                    lat_values = [float(val.num) / float(val.den) for val in latitude.values]
                    long_values = [float(val.num) / float(val.den) for val in longitude.values]

                    # Conversion en degrés décimaux
                    lat = Forensic._convert_to_degrees(lat_values)
                    long = Forensic._convert_to_degrees(long_values)

                    # Appliquer les références N/S et E/W
                    if latitude_ref.values != 'N':
                        lat = 0 -lat
                    if longitude_ref.values != 'E':
                        long = 0 -long

                    print(f"Latitude: {lat}, Longitude: {long}")
                    print("http://maps.google.com/maps?q=loc:%s,%s" % (str(lat), str(long)))
                    if altitude and altitude_ref:
                        alt_ = altitude.values[0]
                        alt = alt_.num / alt_.den
                        if altitude_ref.values[0] == 1:
                            alt = 0 - alt
                        print("ALTITUDE " + str(alt))


    @staticmethod
    def get_pdf_meta(pdf_filename):
        try:
            metadata = []
            with open(pdf_filename, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pdf_doc = pdf_reader.metadata
                print("PDF Metadata:")
                for key, value in pdf_doc.items():
                    print(f"[+] {key}: {value}")
                    metadata.append(f"{key}: {value}")

                # Check for encryption
                if pdf_reader.is_encrypted:
                    print("[!] This PDF is encrypted.")
                else:
                    print("[+] This PDF is not encrypted.")

                # Extract embedded files and objects
                for page in pdf_reader.pages:
                    if "/Annots" in page:
                        print("[+] Found annotations or embedded objects in the PDF.")
                return metadata
        except Exception as e:
            print(f"Error reading PDF metadata: {e}")
            return None


    @staticmethod
    def get_pdf_text(pdf_filename):
        try:
            with open(pdf_filename, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                print("\nPDF Content:")
                for page in pdf_reader.pages:
                    content = page.extract_text()
                    if content:
                        print(content)
                        full_text += content

                # Language Detection
                language = detect(full_text)
                print(f"[+] Detected Language: {language}")

                # Geo Data Extraction
                Forensic.extract_geo_data(full_text)
        except Exception as e:
            print(f"Error reading PDF text: {e}")


    @staticmethod
    def analyze_pdf(pdf_filename):
        metadata = Forensic.get_pdf_meta(pdf_filename)
        if metadata:
            print("\nAnalyzing PDF text...")
            Forensic.get_pdf_text(pdf_filename)
        else:
            print("[-] No metadata found in the PDF.")


    @staticmethod
    def get_docx_meta(docx_filename):
        try:
            doc = Document(docx_filename)
            print("\nDOCX Metadata:")
            core_props = doc.core_properties
            print(f"[+] Title: {core_props.title}")
            print(f"[+] Author: {core_props.author}")
            print(f"[+] Subject: {core_props.subject}")
            print(f"[+] Keywords: {core_props.keywords}")
            print(f"[+] Created: {core_props.created}")

            # Analyze revisions and comments
            for rel in doc.part.rels.values():
                if "comments" in rel.reltype:
                    print("[+] Document contains comments.")
                if "revisions" in rel.reltype:
                    print("[+] Document contains revisions.")
        except Exception as e:
            print(f"Error reading DOCX metadata: {e}")


    @staticmethod
    def get_docx_text(docx_filename):
        try:
            doc = Document(docx_filename)
            full_text = ""
            print("\nDOCX Content:")
            for para in doc.paragraphs:
                print(para.text)
                full_text += para.text

            # Language Detection
            language = detect(full_text)
            print(f"[+] Detected Language: {language}")

            # Geo Data Extraction
            Forensic.extract_geo_data(full_text)
        except Exception as e:
            print(f"Error reading DOCX text: {e}")


    @staticmethod
    def get_excel_meta(excel_filename):
        try:
            xls = pd.ExcelFile(excel_filename)
            print("\nExcel Metadata:")
            print(f"[+] Sheet Names: {xls.sheet_names}")

            workbook = openpyxl.load_workbook(excel_filename, data_only=True)
            print(f"[+] Excel Creator: {workbook.properties.creator}")
            print(f"[+] Last Modified By: {workbook.properties.lastModifiedBy}")
        except Exception as e:
            print(f"Error reading Excel metadata: {e}")


    def get_exif(img_filename):

        with open(img_filename, "rb") as file:
            exif = exifread.process_file(file)
            if not exif:
                print("NOTHING METADATA FOUND")
            else:
                for key, value in exif.items():
                    print(key + " : " + str(value))


    @staticmethod
    def get_excel_text(excel_filename):
        try:
            df = pd.read_excel(excel_filename, sheet_name=None)
            full_text = ""
            print("\nExcel Content:")
            for sheet_name, sheet_data in df.items():
                print(f"\nSheet Name: {sheet_name}")
                print(sheet_data.to_string())
                full_text += sheet_data.to_string()

            # Language Detection
            language = detect(full_text)
            print(f"[+] Detected Language: {language}")

            # Geo Data Extraction
            Forensic.extract_geo_data(full_text)

            workbook = openpyxl.load_workbook(excel_filename, data_only=False)
            for sheet in workbook.sheetnames:
                ws = workbook[sheet]
                if ws.sheet_state == 'hidden':
                    print(f"[!] Hidden sheet detected: {sheet}")
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.data_type == 'f':
                            print(f"[+] Formula found in {cell.coordinate}: {cell.value}")
        except Exception as e:
            print(f"Error reading Excel text: {e}")


    @staticmethod
    def extract_geo_data(text):
        try:
            geolocator = Nominatim(user_agent="forensic_analysis")
            coordinates = re.findall(r"\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b", text)
            gps_coords = re.findall(r"\b\d{1,2}\.\d{1,6},\s?\d{1,2}\.\d{1,6}\b", text)

            for coord in coordinates:
                location = geolocator.geocode(coord)
                if location:
                    print(f"[+] GeoData for {coord}: {location.address}")

            for gps in gps_coords:
                latitude, longitude = map(float, gps.split(','))
                location = geolocator.reverse((latitude, longitude))
                if location:
                    print(f"[+] GeoData for GPS {gps}: {location.address}")
        except Exception as e:
            print(f"Error extracting geo data: {e}")


    @staticmethod
    def get_chrome_history(places_db):
        try:
            conn = sqlite3.connect(places_db)  # Use 'conn' for consistency
            cursor = conn.cursor()

            # Improved query for clarity and potential performance gain
            query = """
            SELECT url, datetime(last_visit_date/1000000, "unixepoch") AS visit_date
            FROM moz_places
            INNER JOIN moz_historyvisits ON moz_places.id = moz_historyvisits.place_id
            WHERE visit_count > 0
            """
            cursor.execute(query)

            with open("C:/Users/HP ELITEBOOK/OneDrive/Documents/TutoToutApprendre/rapport_firefox_historique.html", "a") as f:
                # Improved CSS with basic styling
                css = """
                table, th, td {
                    border: 1px solid black;
                    border-collapse: collapse;
                    padding: 5px;
                }
                th {
                    text-align: left;
                }
                """
                header = f"""<!DOCTYPE html>
    <head>
    <style>
    {css}
    </style>
    </head>
    <body>
    <table>
    <tr>
    <th>URL</th>
    <th>Date</th>
    </tr>
    """
                f.write(header)

                for row in cursor:
                    url, visit_date = row  # Unpack row for readability
                    row_html = f"<tr><td><a href='{url}'>{url}</a></td><td>{visit_date}</td></tr>"
                    f.write(row_html)

                footer = "</table></body></html>"
                f.write(footer)

        except Exception as e:
            print("[-] ERROR:", e)  # Use f-string for cleaner error message
            exit(1)

        finally:  # Added a finally block to ensure the connection is closed
            conn.close()

    def get_firefox_cookies(cookies_sqlite):
        """
        Cette fonction extrait les cookies de la base de données SQLite de Firefox et les enregistre dans un fichier HTML.

        Args:
            cookies_sqlite: Le chemin vers la base de données SQLite des cookies Firefox.
        """

        try:
            # Connexion à la base de données SQLite
            conn = sqlite3.connect(cookies_sqlite)
            cursor = conn.cursor()

            # Exécution de la requête SQL pour récupérer les informations des cookies
            cursor.execute("SELECT name, value, host FROM moz_cookies")

            # Création de l'en-tête HTML avec un style simple
            header = """<!DOCTYPE html>
            <head>
            <style>
            table, th, td {
                border: 1px solid blue;
            }
            </style>
            </head>
            <body>
            <table>
            <tr>
                <th>Nom</th>
                <th>Valeur</th>
                <th>Hôte</th>
            </tr>
            """

            # Ouverture du fichier HTML en mode écriture (append)
            with open("C:/Users/HP ELITEBOOK/OneDrive/Documents/TutoToutApprendre/rapport_firefox_cookies.html", "a") as f:
                f.write(header)

                # Itération sur les résultats de la requête
                for row in cursor:
                    name, value, host = row
                    row_html = f"<tr><td>{name}</td><td>{value}</td><td>{host}</td></tr>"
                    f.write(row_html)

                # Fermeture de la table HTML
                footer = "</table></body></html>"
                f.write(footer)

        except Exception as e:
            print(f"[-] Erreur: {str(e)}")
            exit(1)
        finally:
            # Fermeture de la connexion à la base de données
            conn.close()

